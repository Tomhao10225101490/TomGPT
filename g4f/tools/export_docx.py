"""Generate downloadable .docx files from markdown/plain text."""

from __future__ import annotations

import io
import re
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING

    has_docx = True
except ImportError:
    has_docx = False


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def build_docx_bytes(
    content: str,
    title: Optional[str] = None,
) -> bytes:
    if not has_docx:
        raise RuntimeError('Install "python-docx" | pip install python-docx')

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    if title and str(title).strip():
        heading = doc.add_heading(_strip_md_inline(str(title)), level=1)
        for run in heading.runs:
            run.font.size = Pt(18)

    # Normalize newlines; keep fenced code blocks as plain paragraphs
    text = (content or "").replace("\r\n", "\n").strip()
    if not text:
        doc.add_paragraph("(empty)")
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # Remove simple fenced code markers but keep inner text
    text = re.sub(r"```[\w.-]*\n?", "", text)
    text = text.replace("```", "")

    for raw_line in text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(_strip_md_inline(heading_match.group(2)), level=level)
            continue

        bullet_match = re.match(r"^[-*+]\s+(.*)$", line.strip())
        if bullet_match:
            p = doc.add_paragraph(_strip_md_inline(bullet_match.group(1)), style="List Bullet")
            _set_paragraph_spacing(p)
            continue

        numbered_match = re.match(r"^(\d+)[.)]\s+(.*)$", line.strip())
        if numbered_match:
            p = doc.add_paragraph(_strip_md_inline(numbered_match.group(2)), style="List Number")
            _set_paragraph_spacing(p)
            continue

        p = doc.add_paragraph(_strip_md_inline(line))
        _set_paragraph_spacing(p)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_paragraph_spacing(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def safe_docx_filename(name: Optional[str], fallback: str = "TomGPT.docx") -> str:
    raw = (name or fallback).strip() or fallback
    raw = re.sub(r"[\\/:*?\"<>|]+", "-", raw)
    if not raw.lower().endswith(".docx"):
        raw = f"{raw}.docx"
    return raw[:180]
